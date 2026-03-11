"""
Convert All_Participants_Timetable.md to a Word document (.docx).
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def md_table_to_rows(line_iter):
    """Yield rows (list of cell strings) from a markdown table. line_iter advances."""
    rows = []
    for line in line_iter:
        line = line.rstrip()
        if not line.startswith("|"):
            break
        # Remove leading/trailing |
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(c.replace("-", "") == "" for c in cells):
            continue  # skip separator row
        rows.append(cells)
    return rows

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    doc.add_heading("All Participants – Timetables (PathFAInder)", 0)
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("|"):
            class Iter:
                def __init__(self, idx):
                    self.idx = idx
                def __iter__(self):
                    return self
                def __next__(self):
                    if self.idx >= len(lines):
                        raise StopIteration
                    out = lines[self.idx]
                    self.idx += 1
                    return out
            it = Iter(i)
            rows = md_table_to_rows(it)
            i = it.idx - 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
        elif stripped == "---" or stripped == "":
            pass
        elif stripped.startswith("*") and stripped.endswith("*"):
            p = doc.add_paragraph(stripped.strip("*"))
            p.paragraph_format.space_after = Pt(6)
        else:
            doc.add_paragraph(stripped)
        i += 1
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx("All_Participants_Timetable.md", "All_Participants_Timetable.docx")
