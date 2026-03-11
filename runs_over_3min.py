"""
List all Follow path runs that last more than 3 minutes. Output: Word doc and print.
"""
import os
import json
import sqlite3
import re
from datetime import datetime

from docx import Document

BASE = "Participants-data-pathFinder/Participants-data-pathFinder"
MIN_DURATION_MIN = 3
OUT_DOCX = "Runs_Over_3_Minutes.docx"


def parse_nav_value(value):
    if not value:
        return None
    try:
        j = json.loads(value)
        route = j.get("route") or ""
        args_raw = j.get("args")
        args = json.loads(args_raw) if isinstance(args_raw, str) else None
        path_uri = None
        with_haptic = None
        if args and isinstance(args, dict) and args.get("mMap"):
            path_uri = args["mMap"].get("pathUri")
            with_haptic = args["mMap"].get("withHaptic")
        return {"route": route, "pathUri": path_uri, "withHaptic": with_haptic}
    except (json.JSONDecodeError, TypeError):
        return None


def route_to_activity(route, path_uri, with_haptic):
    if "FollowThePathScreenType" not in route:
        return None
    mode = "Haptic" if with_haptic in (True, 1, 2) else "Visual"
    path_name = os.path.basename(path_uri) if path_uri else "?"
    return f"Follow path ({path_name}) – {mode}"


def folder_to_env(folder_name):
    lower = folder_name.lower()
    if "non" in lower and "urban" in lower:
        return "Non-Urban"
    return "Urban"


def get_last_location_timestamp_ms(db_path, after_ms):
    if not os.path.isfile(db_path):
        return None
    try:
        conn = sqlite3.connect(db_path)
        cur = conn.cursor()
        cur.execute(
            "SELECT MAX(timestamp) FROM logs WHERE type IN ('location', 'raw-location') AND timestamp >= ?",
            (after_ms,),
        )
        row = cur.fetchone()
        conn.close()
        return row[0] if row and row[0] else None
    except Exception:
        return None


def get_participant_dbs(participant_dir):
    out = []
    p_path = os.path.join(BASE, participant_dir)
    if not os.path.isdir(p_path):
        return out
    for type_dir in sorted(os.listdir(p_path)):
        t_path = os.path.join(p_path, type_dir)
        if not os.path.isdir(t_path):
            continue
        for f in os.listdir(t_path):
            if f.lower().endswith(".db") and "wrong" not in f.lower():
                out.append((type_dir, os.path.join(t_path, f)))
                break
    return out


def participant_sort_key(name):
    m = re.search(r"P(\d+)", name, re.I)
    return (int(m.group(1)), name) if m else (999, name)


def get_runs_over_3min():
    """Yield (participant_id, when_str, what_str, duration_min) for runs > MIN_DURATION_MIN. Deduplicated by (p_id, start_ts, path, mode)."""
    participant_dirs = sorted(
        [d for d in os.listdir(BASE) if os.path.isdir(os.path.join(BASE, d)) and re.match(r"P\d+", d, re.I)],
        key=participant_sort_key,
    )
    seen = set()  # (p_id, ts, path_name, mode) to avoid duplicate runs from multiple DBs
    for p_id in participant_dirs:
        dbs = get_participant_dbs(p_id)
        folder_to_path = {f: path for f, path in dbs}
        all_events = []  # (ts, folder, activity)
        for folder, db_path in dbs:
            if not os.path.isfile(db_path):
                continue
            try:
                conn = sqlite3.connect(db_path)
                cur = conn.cursor()
                cur.execute("SELECT timestamp, value FROM logs WHERE type='navigation' ORDER BY timestamp")
                for ts, val in cur.fetchall():
                    data = parse_nav_value(val)
                    if not data:
                        continue
                    activity = route_to_activity(
                        data.get("route"), data.get("pathUri"), data.get("withHaptic")
                    )
                    if activity and "Follow path (" in activity:
                        all_events.append((ts, folder, activity))
                conn.close()
            except Exception:
                continue
        all_events.sort(key=lambda x: x[0])
        all_ts = sorted(set(e[0] for e in all_events))

        def next_ts_after(t):
            for t2 in all_ts:
                if t2 > t:
                    return t2
            return None

        for ts, folder, activity in all_events:
            rest = activity.split("Follow path (", 1)[1]
            if ")" not in rest:
                continue
            path_name = rest.split(")", 1)[0].strip()
            mode = rest.split(")", 1)[1].strip().lstrip(" –-").strip() or "?"
            env = folder_to_env(folder)
            what = f"{env} – {path_name} ({mode})"

            next_ts = next_ts_after(ts)
            if next_ts is None:
                db_path = folder_to_path.get(folder)
                next_ts = get_last_location_timestamp_ms(db_path, ts) if db_path else None
            if next_ts is None:
                continue
            duration_min = (next_ts - ts) / 60000.0
            if duration_min < MIN_DURATION_MIN:
                continue
            key = (p_id, ts, path_name, mode)
            if key in seen:
                continue
            seen.add(key)
            start_dt = datetime.utcfromtimestamp(ts / 1000)
            end_dt = datetime.utcfromtimestamp(next_ts / 1000)
            when = f"{start_dt.strftime('%H:%M')}–{end_dt.strftime('%H:%M')}"
            yield (p_id, when, what, round(duration_min, 1))


def main():
    if not os.path.isdir(BASE):
        print(f"Base not found: {BASE}")
        return
    rows = list(get_runs_over_3min())
    print(f"Runs longer than {MIN_DURATION_MIN} minutes: {len(rows)}")
    print()
    print("#\tParticipant\tWhen (UTC)\tWhat\tDuration (min)")
    print("-" * 80)
    for i, (p_id, when, what, dur) in enumerate(rows, 1):
        print(f"{i}\t{p_id}\t{when}\t{what}\t{dur}")

    doc = Document()
    doc.add_heading(f"Runs longer than {MIN_DURATION_MIN} minutes", 0)
    doc.add_paragraph(f"All times in UTC. Total: {len(rows)} runs.")
    doc.add_paragraph()
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.rows[0].cells[0].text = "#"
    table.rows[0].cells[1].text = "Participant"
    table.rows[0].cells[2].text = "When (UTC)"
    table.rows[0].cells[3].text = "What"
    table.rows[0].cells[4].text = "Duration (min)"
    for i, (p_id, when, what, dur) in enumerate(rows, 1):
        table.rows[i].cells[0].text = str(i)
        table.rows[i].cells[1].text = p_id
        table.rows[i].cells[2].text = when
        table.rows[i].cells[3].text = what
        table.rows[i].cells[4].text = str(dur)
    try:
        doc.save(OUT_DOCX)
        print(f"\nWritten: {OUT_DOCX}")
    except PermissionError:
        doc.save("Runs_Over_3_Minutes_new.docx")
        print("\nOriginal in use. Written: Runs_Over_3_Minutes_new.docx")


if __name__ == "__main__":
    main()
