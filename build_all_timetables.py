"""
Build timetables for ALL participants (P1–P14) and write one Word document (.docx).
"""
import os
import json
import sqlite3
import re
from datetime import datetime

from docx import Document

BASE = "Participants-data-pathFinder/Participants-data-pathFinder"
DOCX_PATH = "All_Participants_Timetable.docx"  # Close in Word before re-running


def parse_nav_value(value):
    if not value:
        return None
    try:
        j = json.loads(value)
        route = j.get("route") or ""
        args_raw = j.get("args")
        args = json.loads(args_raw) if isinstance(args_raw, str) else None
        path_uri = None
        with_haptic = None
        if args and isinstance(args, dict) and args.get("mMap"):
            path_uri = args["mMap"].get("pathUri")
            with_haptic = args["mMap"].get("withHaptic")
        return {"route": route, "pathUri": path_uri, "withHaptic": with_haptic}
    except (json.JSONDecodeError, TypeError):
        return None


def route_to_activity(route, path_uri, with_haptic):
    if "Home" in route:
        return "Home"
    if "NavigationTypeSelectScreenType" in route:
        return "Navigation type select (choose path)"
    if "FollowThePathScreenType" in route:
        mode = "Haptic" if with_haptic in (True, 1, 2) else "Visual"
        path_name = os.path.basename(path_uri) if path_uri else "?"
        return f"Follow path ({path_name}) – {mode}"
    return route.split(".")[-1][:40] if route else "?"


def folder_to_env(folder_name):
    """Urban or Non-Urban from folder e.g. P1-Non-Urban-Haptic -> Non-Urban."""
    lower = folder_name.lower()
    if "non" in lower and "urban" in lower:
        return "Non-Urban"
    return "Urban"


def build_summary_rows(events, by_folder):
    """
    Build list of (#, when_str, what_str) for Follow path only.
    When = "HH:MM–HH:MM" or "HH:MM (one time)". What = "Urban – path.gpx (Visual)" etc.
    """
    all_ts = sorted(set(e[0] for e in events))
    follow_entries = []  # (ts, folder, path_name, mode)
    for ts, folder, activity in events:
        if "Follow path (" not in activity:
            continue
        rest = activity.split("Follow path (", 1)[1]
        if ")" not in rest:
            continue
        path_name = rest.split(")", 1)[0].strip()
        mode = rest.split(")", 1)[1].strip().lstrip(" –-").strip() or "?"
        follow_entries.append((ts, folder, path_name, mode))

    def next_ts_after(t):
        for t2 in all_ts:
            if t2 > t:
                return t2
        return None

    run_num_per_key = {}  # (folder, path_name, mode) -> 1, 2, 3...
    summary = []
    for idx, (ts, folder, path_name, mode) in enumerate(follow_entries):
        key = (folder, path_name, mode)
        run_num_per_key[key] = run_num_per_key.get(key, 0) + 1
        run_n = run_num_per_key[key]
        env = folder_to_env(folder)
        what = f"{env} – {path_name} ({mode})"
        if run_n > 1:
            what += f", {run_n}th run" if run_n > 2 else ", 2nd run"
        else:
            # Check if this key appears again later -> add "1st run"
            if sum(1 for e in follow_entries[idx + 1:] if (e[1], e[2], e[3]) == key) > 0:
                what += ", 1st run"
        next_ts = next_ts_after(ts)
        start_dt = datetime.utcfromtimestamp(ts / 1000)
        if next_ts is None:
            when = f"{start_dt.strftime('%H:%M')} (one time)"
        else:
            end_dt = datetime.utcfromtimestamp(next_ts / 1000)
            when = f"{start_dt.strftime('%H:%M')}–{end_dt.strftime('%H:%M')}"
        summary.append((len(summary) + 1, when, what))
    return summary


def get_participant_dbs(participant_dir):
    """Return list of (folder_name, db_path) for a participant e.g. P1."""
    out = []
    p_path = os.path.join(BASE, participant_dir)
    if not os.path.isdir(p_path):
        return out
    for type_dir in sorted(os.listdir(p_path)):
        t_path = os.path.join(p_path, type_dir)
        if not os.path.isdir(t_path):
            continue
        for f in os.listdir(t_path):
            if f.lower().endswith(".db") and "wrong" not in f.lower():
                out.append((type_dir, os.path.join(t_path, f)))
                break
    return out


def build_timetable_for_participant(participant_id):
    """Return (events list, list of (folder, events) for this participant)."""
    p_dir = participant_id  # e.g. P1
    dbs = get_participant_dbs(p_dir)
    events = []
    by_folder = []
    for folder, db_path in dbs:
        if not os.path.isfile(db_path):
            continue
        try:
            conn = sqlite3.connect(db_path)
            cur = conn.cursor()
            cur.execute(
                "SELECT timestamp, value FROM logs WHERE type='navigation' ORDER BY timestamp"
            )
            rows = cur.fetchall()
            conn.close()
        except Exception:
            continue
        folder_events = []
        for ts, val in rows:
            data = parse_nav_value(val)
            if not data:
                continue
            activity = route_to_activity(
                data.get("route"), data.get("pathUri"), data.get("withHaptic")
            )
            item = (ts, folder, activity)
            events.append(item)
            folder_events.append(item)
        if folder_events:
            by_folder.append((folder, folder_events))
    events.sort(key=lambda x: x[0])
    return events, by_folder


def participant_sort_key(name):
    m = re.search(r"P(\d+)", name, re.I)
    return (int(m.group(1)), name) if m else (999, name)


def main():
    if not os.path.isdir(BASE):
        print(f"Base not found: {BASE}")
        return
    participant_dirs = sorted(
        [d for d in os.listdir(BASE) if os.path.isdir(os.path.join(BASE, d)) and re.match(r"P\d+", d, re.I)],
        key=participant_sort_key,
    )
    doc = Document()
    doc.add_heading("All Participants – Timetables (PathFAInder)", 0)
    doc.add_paragraph("All times in UTC. Summary of Follow path runs only.")
    doc.add_paragraph()

    for p_id in participant_dirs:
        events, by_folder = build_timetable_for_participant(p_id)
        doc.add_heading(p_id, level=1)
        summary_rows = build_summary_rows(events, by_folder)
        if not summary_rows:
            doc.add_paragraph("No Follow path data.")
            doc.add_paragraph()
            continue
        table = doc.add_table(rows=1 + len(summary_rows), cols=3)
        table.rows[0].cells[0].text = "#"
        table.rows[0].cells[1].text = "When (UTC)"
        table.rows[0].cells[2].text = "What"
        for i, (num, when, what) in enumerate(summary_rows, 1):
            table.rows[i].cells[0].text = str(num)
            table.rows[i].cells[1].text = when
            table.rows[i].cells[2].text = what
        doc.add_paragraph()

    try:
        doc.save(DOCX_PATH)
    except PermissionError:
        alt = "All_Participants_Timetable_new.docx"
        doc.save(alt)
        print(f"Original file in use. Written: {alt}")
        return
    print(f"Written: {DOCX_PATH}")


if __name__ == "__main__":
    main()
